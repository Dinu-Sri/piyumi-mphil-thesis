#!/usr/bin/env python
"""Build a versioned Word thesis draft from Markdown chapter files."""

from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONFIG_PATH = ROOT / "tools" / "thesis_build_config.json"
OUTPUT_DIR = ROOT / "output" / "word"
MANIFEST_PATH = OUTPUT_DIR / "build_manifest.jsonl"


def load_config(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_toc_entries(config: dict) -> list[dict]:
    entries_path = config.get("toc_entries_path")
    if not entries_path:
        return []
    path = Path(entries_path)
    if not path.is_absolute():
        path = ROOT / path
    if not path.exists():
        raise FileNotFoundError(f"Configured TOC entries file not found: {path}")
    entries = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(entries, list):
        raise ValueError(f"TOC entries file must contain a JSON list: {path}")
    return entries


def git_commit() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=ROOT,
            text=True,
            capture_output=True,
            check=True,
        )
        return result.stdout.strip()
    except Exception:
        return "unknown"


def next_version(prefix: str) -> tuple[int, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    existing = sorted(OUTPUT_DIR.glob(f"{prefix}_v*.docx"))
    max_version = 0
    for path in existing:
        match = re.search(r"_v(\d+)_", path.name)
        if match:
            max_version = max(max_version, int(match.group(1)))
    version = max_version + 1
    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M")
    return version, OUTPUT_DIR / f"{prefix}_v{version:03d}_{stamp}.docx"


def set_font(run, name: str, size_pt: int | float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document, line_spacing: float) -> None:
    section = doc.sections[0]
    configure_section_layout(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal_rpr = normal._element.get_or_add_rPr()
    normal_fonts = normal_rpr.rFonts
    if normal_fonts is None:
        normal_fonts = OxmlElement("w:rFonts")
        normal_rpr.append(normal_fonts)
    normal_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_fonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_after = Pt(0)

    for style in styles:
        if hasattr(style, "font") and style.font is not None:
            try:
                style.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass

    for name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style_rpr = style._element.get_or_add_rPr()
        style_fonts = style_rpr.rFonts
        if style_fonts is None:
            style_fonts = OxmlElement("w:rFonts")
            style_rpr.append(style_fonts)
        style_fonts.set(qn("w:ascii"), "Times New Roman")
        style_fonts.set(qn("w:hAnsi"), "Times New Roman")
        if name.startswith("Heading"):
            style.font.bold = True
            style.font.size = Pt(14 if name == "Heading 1" else 12)
        else:
            style.font.bold = False
            style.font.size = Pt(12)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 6 if name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(6 if name.startswith("Heading") else 0)


def configure_section_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.7)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.footer_distance = Cm(1.0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font(run, "Times New Roman", 12)


def add_title_page(doc: Document, config: dict) -> None:
    title = config.get("title", "THESIS TITLE")
    author = config.get("author", "FULL NAME OF STUDENT")
    degree = config.get("degree", "MASTER OF PHILOSOPHY")
    university = config.get("university", "WAYAMBA UNIVERSITY OF SRI LANKA")
    year = str(dt.date.today().year)

    for text, size, bold in [
        (title.upper(), 14, True),
        ("", 12, False),
        ("By", 12, False),
        (author.upper(), 12, True),
        ("", 12, False),
        (f"Thesis submitted to the {university} in fulfilment of the requirements for the Degree of {degree}", 12, False),
        ("", 12, False),
        (dt.datetime.now().strftime("%B %Y"), 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        if text:
            run = p.add_run(text)
            set_font(run, "Times New Roman", size, bold=bold)

def set_page_numbering(section, start: int = 1, fmt: str | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)


def add_static_toc_entries(doc: Document, entries: list[dict]) -> None:
    for entry in entries:
        text = str(entry.get("text", "")).strip()
        page = entry.get("thesis_page", entry.get("page"))
        level = int(entry.get("level", 1))
        if not text or page is None:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.5 if level == 2 else 1.0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        title_run = p.add_run(text.upper() if level == 1 else text)
        set_font(title_run, "Times New Roman", 12, bold=(level == 1))
        tab_run = p.add_run("\t")
        set_font(tab_run, "Times New Roman", 12)
        page_run = p.add_run(str(page))
        set_font(page_run, "Times New Roman", 12)


def add_live_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update fields in Word if page numbers are not shown."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)
    set_font(run, "Times New Roman", 12)


def add_table_of_contents(doc: Document, entries: list[dict] | None = None) -> None:
    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(toc_section)
    set_page_numbering(toc_section, 1, "lowerRoman")
    toc_section.footer.is_linked_to_previous = False
    if toc_section.footer.paragraphs:
        add_page_number(toc_section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run("TABLE OF CONTENTS")
    set_font(run, "Times New Roman", 14, bold=True)

    if entries:
        add_static_toc_entries(doc, entries)
    else:
        add_live_toc_field(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(body_section)
    set_page_numbering(body_section, 1)
    body_section.footer.is_linked_to_previous = False
    if body_section.footer.paragraphs:
        add_page_number(body_section.footer.paragraphs[0])


def add_body_section_without_toc(doc: Document) -> None:
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(body_section)
    set_page_numbering(body_section, 1)
    body_section.footer.is_linked_to_previous = False
    if body_section.footer.paragraphs:
        add_page_number(body_section.footer.paragraphs[0])


def strip_front_matter(markdown: str) -> str:
    text = markdown.lstrip("\ufeff")
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            text = text[end + 4 :]
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    return text.strip()


def add_runs_with_inline_formatting(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_font(run, "Times New Roman", 12)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, "Times New Roman", 12, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Times New Roman", 12, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, "Times New Roman", 12)


def add_markdown_table(doc: Document, table_lines: list[str], line_spacing: float) -> None:
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        cells = [cell.strip().replace("\\|", "|") for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = line_spacing
            add_runs_with_inline_formatting(p, text)
            for run in p.runs:
                set_font(run, "Times New Roman", 11, bold=(r_idx == 0))


def add_image(doc: Document, line: str, source_md: Path) -> bool:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return False
    caption, raw_path = match.groups()
    image_path = Path(raw_path.strip())
    if not image_path.is_absolute():
        image_path = (source_md.parent / image_path).resolve()
        if not image_path.exists():
            image_path = (ROOT / raw_path.strip()).resolve()
    if not image_path.exists():
        p = doc.add_paragraph()
        add_runs_with_inline_formatting(p, f"[Missing image: {raw_path}]")
        return True
    doc.add_picture(str(image_path), width=Cm(13))
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(caption)
        set_font(run, "Times New Roman", 11)
    return True


def add_markdown_file(doc: Document, path: Path, line_spacing: float, first_chapter: bool) -> None:
    markdown = strip_front_matter(path.read_text(encoding="utf-8"))
    lines = markdown.splitlines()
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_paragraph():
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = line_spacing
            add_runs_with_inline_formatting(p, text)
            paragraph_buffer.clear()

    def flush_table():
        if table_buffer:
            add_markdown_table(doc, table_buffer, line_spacing)
            table_buffer.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            flush_table()
            continue

        if line.startswith("|"):
            flush_paragraph()
            table_buffer.append(line)
            continue
        flush_table()

        if add_image(doc, line, path):
            flush_paragraph()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                if not first_chapter:
                    doc.add_page_break()
                p = doc.add_paragraph(style="Heading 1")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text.upper())
                set_font(run, "Times New Roman", 14, bold=True)
            elif level == 2:
                p = doc.add_paragraph(style="Heading 2")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                set_font(run, "Times New Roman", 12, bold=True)
            else:
                p = doc.add_paragraph(style="Heading 3")
                p.paragraph_format.left_indent = Cm(1.0)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(text)
                set_font(run, "Times New Roman", 12, bold=True)
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = line_spacing
            add_runs_with_inline_formatting(p, bullet.group(1))
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = line_spacing
            add_runs_with_inline_formatting(p, numbered.group(1))
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    flush_table()


def build_docx(config: dict, stage: str, include_placeholders: bool) -> tuple[Path, dict]:
    line_spacing = 2.0 if stage == "initial" else 1.5
    doc = Document()
    configure_document(doc, line_spacing)

    if config.get("include_title_page", True):
        add_title_page(doc, config)

    toc_mode = config.get("table_of_contents_mode", "static")
    toc_entries = load_toc_entries(config) if toc_mode == "static" else []
    if config.get("include_table_of_contents", True):
        add_table_of_contents(doc, toc_entries)
    else:
        add_body_section_without_toc(doc)

    first = True
    included = []
    skipped = []
    for chapter_rel in config["chapter_order"]:
        path = ROOT / chapter_rel
        if not path.exists():
            skipped.append(chapter_rel)
            continue
        text = path.read_text(encoding="utf-8")
        if not include_placeholders and "placeholder - to be drafted" in text:
            skipped.append(chapter_rel)
            continue
        add_markdown_file(doc, path, line_spacing, first_chapter=first)
        included.append(chapter_rel)
        first = False

    prefix = config.get("output_prefix", "thesis")
    version, output_path = next_version(prefix)
    doc.save(output_path)
    enable_update_fields_on_open(output_path)
    force_ooxml_font_colors_black(output_path)

    manifest = {
        "version": version,
        "created": dt.datetime.now().isoformat(timespec="seconds"),
        "stage": stage,
        "git_commit": git_commit(),
        "output": output_path.relative_to(ROOT).as_posix(),
        "toc_mode": toc_mode if config.get("include_table_of_contents", True) else "none",
        "included": included,
        "skipped": skipped,
    }
    MANIFEST_PATH.parent.mkdir(parents=True, exist_ok=True)
    with MANIFEST_PATH.open("a", encoding="utf-8", newline="\n") as handle:
        handle.write(json.dumps(manifest) + "\n")

    return output_path, manifest


def enable_update_fields_on_open(docx_path: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="docx_update_fields_") as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp_path)

        settings_path = tmp_path / "word" / "settings.xml"
        text = settings_path.read_text(encoding="utf-8", errors="ignore")
        if "<w:updateFields" not in text:
            text = text.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
            settings_path.write_text(text, encoding="utf-8", newline="")

        backup = docx_path.with_suffix(".docx.bak")
        shutil.copy2(docx_path, backup)
        try:
            with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for file_path in tmp_path.rglob("*"):
                    if file_path.is_file():
                        zout.write(file_path, file_path.relative_to(tmp_path).as_posix())
        finally:
            backup.unlink(missing_ok=True)


def force_ooxml_font_colors_black(docx_path: Path) -> None:
    """Normalize all Word font color declarations to black.

    python-docx can leave unused built-in style color values in styles.xml.
    This post-save OOXML pass keeps the thesis submission rule simple: every
    explicit w:color declaration in the generated package is black.
    """

    with tempfile.TemporaryDirectory(prefix="docx_black_font_") as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp_path)

        for xml_path in (tmp_path / "word").rglob("*.xml"):
            text = xml_path.read_text(encoding="utf-8", errors="ignore")

            def repl(match: re.Match[str]) -> str:
                tag = match.group(0)
                tag = re.sub(r'\s+w:themeColor="[^"]*"', "", tag)
                tag = re.sub(r'\s+w:themeTint="[^"]*"', "", tag)
                tag = re.sub(r'\s+w:themeShade="[^"]*"', "", tag)
                if 'w:val="' in tag:
                    tag = re.sub(r'w:val="[^"]*"', 'w:val="000000"', tag)
                else:
                    tag = tag[:-2] + ' w:val="000000"/>'
                return tag

            text = re.sub(r"<w:color\b[^>]*/>", repl, text)
            xml_path.write_text(text, encoding="utf-8", newline="")

        backup = docx_path.with_suffix(".docx.bak")
        shutil.copy2(docx_path, backup)
        try:
            with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for file_path in tmp_path.rglob("*"):
                    if file_path.is_file():
                        zout.write(file_path, file_path.relative_to(tmp_path).as_posix())
        finally:
            backup.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a versioned thesis DOCX from Markdown chapters.")
    parser.add_argument("--config", default=str(CONFIG_PATH), help="Build config JSON path.")
    parser.add_argument("--stage", choices=["initial", "final"], default="final", help="Line-spacing mode.")
    parser.add_argument("--include-placeholders", action="store_true", help="Include placeholder chapter files.")
    args = parser.parse_args()

    config_path = Path(args.config)
    if not config_path.is_absolute():
        config_path = ROOT / config_path
    output_path, manifest = build_docx(load_config(config_path), args.stage, args.include_placeholders)
    print(f"created {output_path.relative_to(ROOT)}")
    print(json.dumps(manifest, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
